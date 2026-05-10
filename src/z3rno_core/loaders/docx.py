"""DOCX loader — python-docx-backed extraction.

Walks a Word document's paragraphs in order and emits plain text with
preserved blank-line boundaries (so the Phase A paragraph chunker has
something to split on). Headings are rendered as Markdown-style
prefixes (``# `` for level 1, ``## `` for level 2, ...) so downstream
summaries retain the document's structure without depending on a
DOCX-aware reader at every step.
"""

from __future__ import annotations

import io
import logging
import zipfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from z3rno_core.loaders.base import (
    Loader,
    LoaderInputError,
    LoaderResult,
)

logger = logging.getLogger(__name__)


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocxLoader(Loader):
    """Extract paragraphs (with heading hints) from a .docx file."""

    @property
    def name(self) -> str:
        return "docx"

    async def load(
        self,
        content: bytes,
        *,
        filename: str | None = None,
        mime_type: str | None = None,
    ) -> LoaderResult:
        if not content:
            return LoaderResult(
                text="",
                metadata={
                    "loader": self.name,
                    "mime_type": _DOCX_MIME,
                    "filename": filename,
                    "byte_size": 0,
                    "char_count": 0,
                    "paragraph_count": 0,
                    "heading_count": 0,
                },
            )

        try:
            document = Document(io.BytesIO(content))
        except (PackageNotFoundError, zipfile.BadZipFile) as exc:
            raise LoaderInputError(f"malformed DOCX: {exc}") from exc

        lines: list[str] = []
        heading_count = 0
        for paragraph in document.paragraphs:
            text = paragraph.text or ""
            if not text.strip():
                # Preserve blank-paragraph boundaries so the chunker can split.
                lines.append("")
                continue

            style_name = (paragraph.style.name if paragraph.style else "") or ""
            level = _heading_level(style_name)
            if level is not None:
                heading_count += 1
                lines.append(f"{'#' * level} {text}")
            else:
                lines.append(text)

        # Join paragraphs with newlines; keep blank paragraphs as the
        # paragraph chunker's natural boundary.
        body = "\n".join(lines).strip()

        return LoaderResult(
            text=body,
            metadata={
                "loader": self.name,
                "mime_type": _DOCX_MIME,
                "filename": filename,
                "byte_size": len(content),
                "char_count": len(body),
                "paragraph_count": len(document.paragraphs),
                "heading_count": heading_count,
            },
        )


def _heading_level(style_name: str) -> int | None:
    """Map a docx paragraph-style name to a Markdown heading level (1..6).

    Returns ``None`` for non-heading styles. Cap at 6 to match Markdown's
    range; Word's "Heading 7+" reduces to 6.
    """
    name = style_name.strip().lower()
    if not name.startswith("heading"):
        return None
    suffix = name[len("heading") :].strip()
    try:
        level = int(suffix)
    except ValueError:
        return None
    return max(1, min(6, level))
