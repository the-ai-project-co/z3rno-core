"""Unit tests for z3rno_core.loaders (Phase B.1).

Pure-function tests; no DB, no LLM, no network. URL-fetch tests mock
``httpx.AsyncClient``.
"""

from __future__ import annotations

import asyncio
import io
from unittest.mock import AsyncMock, MagicMock, patch

import httpx
import pytest
from docx import Document
from pypdf import PdfWriter

from z3rno_core.loaders import (
    CodeLoader,
    CsvLoader,
    DocxLoader,
    Loader,
    LoaderInputError,
    LoaderRegistry,
    LoaderResult,
    MarkdownLoader,
    PdfLoader,
    PlainTextLoader,
    UnsupportedContentTypeError,
    UrlFetchError,
    UrlLoader,
    fetch_url,
    get_default_registry,
    sniff_mime_type,
)

# ---------------------------------------------------------------------------
# Helpers — fixture-on-the-fly so tests don't need files on disk.
# ---------------------------------------------------------------------------


def _build_blank_pdf(num_pages: int) -> bytes:
    w = PdfWriter()
    for _ in range(num_pages):
        w.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    w.write(buf)
    return buf.getvalue()


def _build_docx() -> bytes:
    doc = Document()
    doc.add_heading("Z3rno", level=1)
    doc.add_paragraph("Z3rno is a smart-memory platform.")
    doc.add_heading("Architecture", level=2)
    doc.add_paragraph("Combines a knowledge graph with vector search.")
    doc.add_paragraph("")
    doc.add_paragraph("Final paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# LoaderResult schema
# ---------------------------------------------------------------------------


class TestLoaderResult:
    def test_defaults(self) -> None:
        r = LoaderResult()
        assert r.text == ""
        assert r.metadata == {}
        assert r.is_empty is True
        assert r.char_count == 0

    def test_char_count_falls_back_to_text_length(self) -> None:
        r = LoaderResult(text="hello", metadata={})
        assert r.char_count == 5

    def test_char_count_from_metadata(self) -> None:
        r = LoaderResult(text="hello", metadata={"char_count": 999})
        assert r.char_count == 999


# ---------------------------------------------------------------------------
# sniff_mime_type
# ---------------------------------------------------------------------------


class TestSniffMimeType:
    def test_pdf(self) -> None:
        assert sniff_mime_type(b"%PDF-1.7\nfoo") == "application/pdf"

    def test_docx(self) -> None:
        assert (
            sniff_mime_type(b"PK\x03\x04...word/document.xml...padding")
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_zip_without_docx_marker(self) -> None:
        assert sniff_mime_type(b"PK\x03\x04random data") == "application/zip"

    def test_empty_returns_none(self) -> None:
        assert sniff_mime_type(b"") is None

    def test_random_text_returns_none(self) -> None:
        assert sniff_mime_type(b"hello world") is None


# ---------------------------------------------------------------------------
# LoaderRegistry
# ---------------------------------------------------------------------------


class TestLoaderRegistry:
    def test_registers_by_mime_and_extension(self) -> None:
        reg = LoaderRegistry()
        loader = PlainTextLoader()
        reg.register(loader, mime_types=["text/plain"], extensions=["txt"])
        assert reg.get_loader(b"x", mime_type="text/plain") is loader
        assert reg.get_loader(b"x", filename="note.txt") is loader

    def test_mime_strips_charset_suffix(self) -> None:
        reg = LoaderRegistry()
        loader = PlainTextLoader()
        reg.register(loader, mime_types=["text/plain"])
        assert reg.get_loader(b"x", mime_type="text/plain; charset=utf-8") is loader

    def test_magic_byte_sniff_takes_precedence_over_extension(self) -> None:
        reg = LoaderRegistry()
        text_loader = PlainTextLoader()
        pdf_loader = PdfLoader()
        reg.register(text_loader, extensions=["txt"], is_fallback=True)
        reg.register(pdf_loader, mime_types=["application/pdf"])
        # PDF magic bytes win over .txt extension
        chosen = reg.get_loader(b"%PDF-1.7\n", filename="oops.txt")
        assert chosen is pdf_loader

    def test_falls_back_when_registered(self) -> None:
        reg = LoaderRegistry()
        loader = PlainTextLoader()
        reg.register(loader, is_fallback=True)
        assert reg.get_loader(b"x", mime_type="unknown/type") is loader

    def test_raises_when_no_fallback(self) -> None:
        reg = LoaderRegistry()
        with pytest.raises(UnsupportedContentTypeError):
            reg.get_loader(b"x", mime_type="unknown")

    def test_known_introspection(self) -> None:
        reg = LoaderRegistry()
        reg.register(PlainTextLoader(), mime_types=["text/plain"], extensions=["txt"])
        assert "text/plain" in reg.known_mime_types
        assert "txt" in reg.known_extensions


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


class TestPlainTextLoader:
    def test_round_trip(self) -> None:
        r = asyncio.run(PlainTextLoader().load(b"hello\nworld\n", filename="x.txt"))
        assert r.text == "hello\nworld\n"
        assert r.metadata["loader"] == "text"
        assert r.metadata["line_count"] == 2

    def test_invalid_utf8_replaces(self) -> None:
        r = asyncio.run(PlainTextLoader().load(b"\xff\xfe ok"))
        assert "ok" in r.text


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


class TestPdfLoader:
    def test_empty_returns_empty_result(self) -> None:
        r = asyncio.run(PdfLoader().load(b""))
        assert r.text == ""
        assert r.metadata["page_count"] == 0

    def test_malformed_raises(self) -> None:
        with pytest.raises(LoaderInputError):
            asyncio.run(PdfLoader().load(b"%PDF-not-really"))

    def test_blank_pages_metadata(self) -> None:
        pdf = _build_blank_pdf(3)
        r = asyncio.run(PdfLoader().load(pdf, filename="blank.pdf"))
        assert r.metadata["page_count"] == 3
        assert r.metadata["mime_type"] == "application/pdf"
        assert r.metadata["filename"] == "blank.pdf"
        assert "page_offsets" in r.metadata


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class TestDocxLoader:
    def test_empty_returns_empty_result(self) -> None:
        r = asyncio.run(DocxLoader().load(b""))
        assert r.text == ""
        assert r.metadata["paragraph_count"] == 0

    def test_malformed_raises(self) -> None:
        with pytest.raises(LoaderInputError):
            asyncio.run(DocxLoader().load(b"not a zip"))

    def test_renders_headings_as_markdown(self) -> None:
        r = asyncio.run(DocxLoader().load(_build_docx(), filename="d.docx"))
        assert "# Z3rno" in r.text
        assert "## Architecture" in r.text
        assert "Final paragraph." in r.text
        assert r.metadata["heading_count"] == 2

    def test_heading_level_helper(self) -> None:
        from z3rno_core.loaders.docx import _heading_level

        assert _heading_level("Heading 1") == 1
        assert _heading_level("Heading 9") == 6  # cap
        assert _heading_level("Heading") is None
        assert _heading_level("Normal") is None


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------


class TestCsvLoader:
    def test_empty_input(self) -> None:
        r = asyncio.run(CsvLoader().load(b""))
        assert r.metadata["row_count"] == 0

    def test_header_only_input(self) -> None:
        r = asyncio.run(CsvLoader().load(b"a,b,c\n"))
        assert r.metadata["row_count"] == 0
        assert r.metadata["column_count"] == 3
        assert r.metadata["truncated"] is False

    def test_basic_csv(self) -> None:
        r = asyncio.run(
            CsvLoader().load(b"name,role\nZ3rno,product\nCognee,product\n", filename="f.csv")
        )
        assert r.metadata["row_count"] == 2
        assert r.metadata["column_count"] == 2
        assert "name: Z3rno; role: product" in r.text

    def test_truncation_respected(self) -> None:
        big = b"a,b\n" + b"x,y\n" * 100
        r = asyncio.run(CsvLoader(max_rows=10).load(big))
        assert r.metadata["row_count"] == 10
        assert r.metadata["truncated"] is True

    def test_max_rows_zero_rejected_at_construction(self) -> None:
        with pytest.raises(ValueError, match="max_rows must be > 0"):
            CsvLoader(max_rows=0)

    def test_bom_tolerated(self) -> None:
        bom = "\ufefffield_a\tfield_b\nx\ty\n".encode()
        r = asyncio.run(CsvLoader().load(bom, filename="data.tsv"))
        assert r.metadata["column_count"] == 2

    def test_padded_short_rows(self) -> None:
        r = asyncio.run(CsvLoader().load(b"a,b,c\nx,y\n"))
        # short row gets padded; doesn't crash
        assert r.metadata["row_count"] == 1


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


class TestMarkdownLoader:
    def test_passthrough_with_counters(self) -> None:
        md = b"# Title\n\n## Sub\n\nParagraph.\n\n```py\nx=1\n```\n"
        r = asyncio.run(MarkdownLoader().load(md, filename="d.md"))
        assert r.metadata["heading_count"] == 2
        assert r.metadata["code_block_count"] == 1

    def test_empty(self) -> None:
        r = asyncio.run(MarkdownLoader().load(b""))
        assert r.text == ""


# ---------------------------------------------------------------------------
# Code
# ---------------------------------------------------------------------------


class TestCodeLoader:
    def test_python_detected(self) -> None:
        r = asyncio.run(CodeLoader().load(b"def f(): pass\n", filename="x.py"))
        assert r.metadata["language"] == "python"

    def test_typescript_detected_for_tsx(self) -> None:
        r = asyncio.run(CodeLoader().load(b"const x = 1\n", filename="x.tsx"))
        assert r.metadata["language"] == "typescript"

    def test_unknown_extension_marked(self) -> None:
        r = asyncio.run(CodeLoader().load(b"data", filename="x.weird"))
        assert r.metadata["language"] == "unknown"

    def test_no_filename_marked_unknown(self) -> None:
        r = asyncio.run(CodeLoader().load(b"data"))
        assert r.metadata["language"] == "unknown"


# ---------------------------------------------------------------------------
# URL — fetch_url + UrlLoader
# ---------------------------------------------------------------------------


class TestFetchUrl:
    def test_rejects_disallowed_scheme(self) -> None:
        with pytest.raises(UrlFetchError, match="scheme"):
            asyncio.run(fetch_url("file:///etc/passwd"))

    def test_rejects_no_host(self) -> None:
        with pytest.raises(UrlFetchError, match="missing host"):
            asyncio.run(fetch_url("http://"))

    def _patch_get(self, fake_resp: object) -> AsyncMock:
        client = AsyncMock()
        client.__aenter__.return_value = client
        client.get = AsyncMock(return_value=fake_resp)
        return client

    def test_success_returns_fetch_result(self) -> None:
        fake = MagicMock(spec=httpx.Response)
        fake.status_code = 200
        fake.url = "https://example.com/x"
        fake.content = b"<html></html>"
        fake.headers = {"content-type": "text/html; charset=utf-8"}
        fake.reason_phrase = "OK"

        client = self._patch_get(fake)
        with patch("z3rno_core.loaders.url.httpx.AsyncClient", return_value=client):
            result = asyncio.run(fetch_url("https://example.com/x"))
        assert result.status_code == 200
        assert result.content_type == "text/html"
        assert b"<html>" in result.content

    def test_404_raises(self) -> None:
        fake = MagicMock(spec=httpx.Response)
        fake.status_code = 404
        fake.url = "https://example.com/missing"
        fake.content = b""
        fake.headers = {"content-type": "text/plain"}
        fake.reason_phrase = "Not Found"

        client = self._patch_get(fake)
        with (
            patch("z3rno_core.loaders.url.httpx.AsyncClient", return_value=client),
            pytest.raises(UrlFetchError, match="404"),
        ):
            asyncio.run(fetch_url("https://example.com/missing"))

    def test_oversize_response_raises(self) -> None:
        fake = MagicMock(spec=httpx.Response)
        fake.status_code = 200
        fake.url = "https://example.com/big"
        fake.content = b"x" * 10_000
        fake.headers = {"content-type": "text/plain"}
        fake.reason_phrase = "OK"

        client = self._patch_get(fake)
        with (
            patch("z3rno_core.loaders.url.httpx.AsyncClient", return_value=client),
            pytest.raises(UrlFetchError, match="exceeds max_bytes"),
        ):
            asyncio.run(fetch_url("https://example.com/big", max_bytes=1024))

    def test_timeout_raises(self) -> None:
        client = AsyncMock()
        client.__aenter__.return_value = client
        client.get = AsyncMock(side_effect=httpx.TimeoutException("slow"))
        with (
            patch("z3rno_core.loaders.url.httpx.AsyncClient", return_value=client),
            pytest.raises(UrlFetchError, match="timeout"),
        ):
            asyncio.run(fetch_url("https://example.com"))


class TestUrlLoader:
    def test_empty_html(self) -> None:
        r = asyncio.run(UrlLoader().load(b"", mime_type="text/html"))
        assert r.text == ""

    def test_html_extraction_drops_noise(self) -> None:
        html = (
            b"<html><head><title>T</title></head>"
            b"<body><nav>nav</nav><article><h1>X</h1>"
            b"<p>main content</p><script>noise</script></article>"
            b"<footer>footer</footer></body></html>"
        )
        r = asyncio.run(UrlLoader().load(html, mime_type="text/html"))
        assert "main content" in r.text
        assert "noise" not in r.text
        assert "nav" not in r.text
        assert "footer" not in r.text
        assert r.metadata["title"] == "T"

    def test_plain_passthrough(self) -> None:
        r = asyncio.run(UrlLoader().load(b"raw text", mime_type="text/plain"))
        assert r.text == "raw text"
        assert r.metadata["mime_type"] == "text/plain"

    def test_markdown_passthrough(self) -> None:
        r = asyncio.run(UrlLoader().load(b"# x", mime_type="text/markdown"))
        assert "# x" in r.text


# ---------------------------------------------------------------------------
# Default registry — confirm every loader is wired in
# ---------------------------------------------------------------------------


class TestDefaultRegistry:
    @pytest.mark.parametrize(
        ("filename", "expected_class"),
        [
            ("note.txt", PlainTextLoader),
            ("doc.pdf", PdfLoader),
            ("note.docx", DocxLoader),
            ("data.csv", CsvLoader),
            ("readme.md", MarkdownLoader),
            ("hello.py", CodeLoader),
            ("page.html", UrlLoader),
        ],
    )
    def test_routes_by_extension(self, filename: str, expected_class: type[Loader]) -> None:
        loader = get_default_registry().get_loader(b"data", filename=filename)
        assert isinstance(loader, expected_class)

    def test_routes_unknown_to_text_fallback(self) -> None:
        loader = get_default_registry().get_loader(b"data", filename="x.weird")
        assert isinstance(loader, PlainTextLoader)
